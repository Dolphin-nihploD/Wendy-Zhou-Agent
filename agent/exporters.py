"""
agent/exporters.py
==================
Document I/O for Wendy — both directions in one module:

  READING an uploaded file into text (so Wendy can read it):
      extract_text(name, media_type, b64) -> str

  WRITING a stored conversation into a real downloadable document:
      conversation_to_docx(title, messages) -> bytes   (Word, python-docx)
      conversation_to_pdf(title, messages)  -> bytes   (PDF,  reportlab)
      conversation_to_md(title, messages)   -> str      (Markdown)

wendy_agent.py imports extract_text from here, and app.py imports the
conversation_to_* helpers from here, so everything lives in one place.

Dependencies (see requirements.txt):
    python-docx   -> .docx  (read + write)
    openpyxl      -> .xlsx  (read)
    reportlab     -> .pdf   (write)
(.csv / .txt / PDF-reading need no library — PDFs go to Claude natively.)
"""

import io
import re
import html
import base64


# =========================================================================
# READING — turn an uploaded, non-image file into text Wendy can read
# =========================================================================

# rough cap so one giant spreadsheet can't blow past the model's context
MAX_CHARS = 200_000


def _clip(text):
    if text and len(text) > MAX_CHARS:
        return text[:MAX_CHARS] + "\n\n…[truncated — file was longer than the read limit]"
    return text


def extract_text(name, media_type, b64):
    """Decode a base64 file and return its text content.

    name       : original filename (used to detect the type by extension)
    media_type : browser-reported MIME type (a secondary hint)
    b64        : base64-encoded file bytes

    Returns a string (never raises — on failure returns a short note so the
    chat can still proceed).
    """
    try:
        raw = base64.b64decode(b64 or "")
    except Exception as e:
        return f"[Could not decode {name or 'file'}: {e}]"

    lname = (name or "").lower()
    mtype = (media_type or "").lower()

    try:
        # ---- Word (.docx) ----
        if lname.endswith(".docx") or "wordprocessingml" in mtype:
            from docx import Document
            doc = Document(io.BytesIO(raw))
            parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            return _clip("\n".join(parts)) or "[The Word document had no readable text.]"

        # ---- Excel (.xlsx) ----
        if lname.endswith(".xlsx") or "spreadsheetml" in mtype:
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
            out = []
            for ws in wb.worksheets:
                out.append(f"## Sheet: {ws.title}")
                for row in ws.iter_rows(values_only=True):
                    cells = ["" if c is None else str(c) for c in row]
                    if any(cell.strip() for cell in cells):
                        out.append(", ".join(cells))
            wb.close()
            return _clip("\n".join(out)) or "[The spreadsheet had no readable cells.]"

        # ---- CSV / plain text / anything text-like ----
        if lname.endswith((".csv", ".txt", ".md", ".json", ".log")) or mtype.startswith("text/"):
            return _clip(raw.decode("utf-8", errors="replace"))

    except Exception as e:
        return f"[Could not read {name or 'file'}: {e}]"

    # last-ditch: try to decode as UTF-8 text
    try:
        return _clip(raw.decode("utf-8", errors="replace"))
    except Exception:
        return f"[Unsupported file type: {name or 'file'}]"


# =========================================================================
# WRITING — turn a stored conversation into a downloadable document
# =========================================================================

def _who(role):
    return "You" if role == "user" else "Wendy"


# --------------------------------------------------------------- Word ----
def conversation_to_docx(title, messages):
    """Build a .docx from the conversation and return its bytes."""
    from docx import Document
    from docx.shared import RGBColor

    doc = Document()
    doc.add_heading(title or "Conversation", level=0)

    for m in messages:
        who = _who(m.get("role"))
        p = doc.add_paragraph()
        run = p.add_run(who)
        run.bold = True
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        body = (m.get("content") or "").strip()
        if body:
            doc.add_paragraph(body)

        art = m.get("artifact")
        if art and art.get("content"):
            cap = doc.add_paragraph()
            crun = cap.add_run((art.get("title") or "Attachment") + ":")
            crun.italic = True
            doc.add_paragraph(art["content"])

        doc.add_paragraph()  # spacer

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------- PDF ----
def _register_unicode_font():
    """Register a Unicode TTF if one is available on the machine.

    Returns the font name to use ('Uni' if a TTF was found, else the built-in
    'Helvetica', which is Latin-1 only). Trying real system fonts first means
    accented characters, em-dashes, and math symbols render on most machines
    without shipping a font file with the app.
    """
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os

    candidates = [
        r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\segoeui.ttf",
        r"C:\Windows\Fonts\calibri.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/Library/Fonts/Arial.ttf",
    ]
    for path in candidates:
        try:
            if os.path.exists(path):
                pdfmetrics.registerFont(TTFont("Uni", path))
                return "Uni"
        except Exception:
            pass
    return "Helvetica"


def conversation_to_pdf(title, messages):
    """Build a PDF from the conversation and return its bytes."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    font_name = _register_unicode_font()
    unicode_ok = font_name == "Uni"

    def clean(s):
        s = s or ""
        # reportlab can't render colour emoji / astral-plane chars — drop them
        s = re.sub(r"[\U00010000-\U0010FFFF]", "", s)
        if not unicode_ok:
            # Helvetica is Latin-1 only; replace anything outside it
            s = s.encode("latin-1", "replace").decode("latin-1")
        s = html.escape(s)
        return s.replace("\n", "<br/>")

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("t", parent=styles["Title"], fontName=font_name)
    who_style = ParagraphStyle(
        "who", parent=styles["Normal"], fontName=font_name,
        fontSize=11, leading=15, spaceBefore=12, spaceAfter=2,
        textColor=colors.HexColor("#333333"),
    )
    body_style = ParagraphStyle(
        "body", parent=styles["Normal"], fontName=font_name,
        fontSize=10, leading=15,
    )
    note_style = ParagraphStyle(
        "note", parent=body_style, textColor=colors.HexColor("#666666"),
    )

    flow = [Paragraph(clean(title or "Conversation"), title_style), Spacer(1, 12)]
    for m in messages:
        flow.append(Paragraph("<b>" + clean(_who(m.get("role"))) + "</b>", who_style))
        body = (m.get("content") or "").strip()
        if body:
            flow.append(Paragraph(clean(body), body_style))
        art = m.get("artifact")
        if art and art.get("content"):
            flow.append(Paragraph("<i>" + clean(art.get("title") or "Attachment") + ":</i>", note_style))
            flow.append(Paragraph(clean(art["content"]), body_style))
        flow.append(Spacer(1, 6))

    buf = io.BytesIO()
    docp = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=0.9 * inch, rightMargin=0.9 * inch,
        topMargin=0.9 * inch, bottomMargin=0.9 * inch,
        title=title or "Conversation",
    )
    docp.build(flow)
    return buf.getvalue()


# ----------------------------------------------------------- Markdown ----
def conversation_to_md(title, messages):
    """Plain-markdown export (kept for parity with the old behaviour)."""
    lines = ["# " + (title or "Conversation"), ""]
    for m in messages:
        lines.append("## " + _who(m.get("role")))
        lines.append((m.get("content") or "").strip())
        art = m.get("artifact")
        if art and art.get("content"):
            lines.append("")
            lines.append("```")
            lines.append(art["content"])
            lines.append("```")
        lines.append("")
    return "\n".join(lines)


def safe_filename(title, default="conversation"):
    name = re.sub(r"[^\w\- ]", "", title or "").strip().replace(" ", "_")
    return name or default


# =========================================================================
# ARTIFACT EXPORT — turn a single piece of content (an artifact Wendy wrote,
# e.g. a document, email, or table) into a real Word / PDF / Excel file.
# =========================================================================

def _docx_add_line(doc, line):
    """Add one line to a docx, honouring simple markdown headings."""
    s = line.rstrip()
    if s.startswith("### "):
        doc.add_heading(s[4:], level=3)
    elif s.startswith("## "):
        doc.add_heading(s[3:], level=2)
    elif s.startswith("# "):
        doc.add_heading(s[2:], level=1)
    else:
        doc.add_paragraph(s)


def text_to_docx(title, content):
    """Build a Word document from arbitrary text/markdown content."""
    from docx import Document

    doc = Document()
    if title:
        doc.add_heading(title, level=0)
    for line in (content or "").split("\n"):
        _docx_add_line(doc, line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def text_to_pdf(title, content):
    """Build a PDF from arbitrary text/markdown content."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    font_name = _register_unicode_font()
    unicode_ok = font_name == "Uni"

    def clean(s):
        s = s or ""
        s = re.sub(r"[\U00010000-\U0010FFFF]", "", s)
        if not unicode_ok:
            s = s.encode("latin-1", "replace").decode("latin-1")
        return html.escape(s)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("t", parent=styles["Title"], fontName=font_name)
    h_style = ParagraphStyle("h", parent=styles["Heading2"], fontName=font_name)
    body_style = ParagraphStyle("b", parent=styles["Normal"], fontName=font_name, fontSize=10, leading=15)

    flow = []
    if title:
        flow.append(Paragraph(clean(title), title_style))
        flow.append(Spacer(1, 10))
    for line in (content or "").split("\n"):
        s = line.rstrip()
        if not s.strip():
            flow.append(Spacer(1, 6))
        elif s.startswith("#"):
            flow.append(Paragraph("<b>" + clean(s.lstrip("# ")) + "</b>", h_style))
        else:
            flow.append(Paragraph(clean(s), body_style))

    buf = io.BytesIO()
    docp = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=0.9 * inch, rightMargin=0.9 * inch,
        topMargin=0.9 * inch, bottomMargin=0.9 * inch,
        title=title or "Document",
    )
    docp.build(flow)
    return buf.getvalue()


def text_to_xlsx(title, content):
    """Build an Excel workbook from arbitrary text. Rows come from lines; cells
    are split on tab, then pipe, then comma (whichever the line uses)."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = (title or "Sheet")[:31] or "Sheet"
    for line in (content or "").split("\n"):
        if line.strip() == "":
            ws.append([])
            continue
        if "\t" in line:
            cells = line.split("\t")
        elif "|" in line:
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
        elif "," in line:
            cells = line.split(",")
        else:
            cells = [line]
        ws.append(cells)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()